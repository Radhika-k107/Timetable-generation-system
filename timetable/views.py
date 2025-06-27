from django.shortcuts import get_object_or_404, render, redirect
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required, user_passes_test
from django.template.defaultfilters import time as time_filter

from .forms import TeacherForm, SubjectForm, ClassroomForm, BatchForm, DivisionForm, YearForm,DepartmentForm, BatchTeacherFormSet
from .models import Teacher, Subject, Classroom, Batch, Division, Year, Timetable, TimeSlot,Department
from .timetable_generator import generate_timetable
import json
from django.http import JsonResponse, HttpResponse
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import qn as qn2
import io
import openpyxl
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph
from reportlab.lib.units import inch
from datetime import datetime
from reportlab.lib.enums import TA_CENTER
from docx.enum.section import WD_ORIENT


def is_super_admin(user):
    return user.is_superuser

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user_type = request.POST.get('user_type', 'admin')
        next_url = request.GET.get('next')

        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            if user_type == 'admin' and user.is_superuser:
                login(request, user)
                return redirect(next_url if next_url else 'main_admin_dashboard')
            elif user_type == 'department':
                try:
                    department = Department.objects.get(admin=user)
                    login(request, user)
                    # If there's a next URL and it's not the admin dashboard, use it
                    if next_url and 'admin-dashboard' not in next_url:
                        return redirect(next_url)
                    return redirect('department_dashboard')
                except Department.DoesNotExist:
                    messages.error(request, 'This user is not associated with any department.')
            else:
                messages.error(request, 'Invalid user type or permissions.')
        else:
            messages.error(request, 'Invalid username or password.')
    
    return render(request, 'login.html', {'user_type': request.POST.get('user_type', 'admin')})


    
@login_required
@user_passes_test(lambda u: u.is_superuser)
def main_admin_dashboard(request):
    departments = Department.objects.all()
    context = {
        'departments': departments,
    }
    return render(request, 'main_admin_dashboard.html', context)

@login_required
def logout_view(request):
    if request.method == 'POST' or request.method == 'GET':
        # Clear all messages
        storage = messages.get_messages(request)
        for message in storage:
            pass
        logout(request)
        return redirect('login')
    
@login_required
@user_passes_test(is_super_admin)
def add_department(request):
    if request.method == 'POST':
        form = DepartmentForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Department added successfully!')
            return redirect('list_departments')
    else:
        form = DepartmentForm()
    departments = Department.objects.all()
    return render(request, 'add_department.html', {'form': form, 'departments': departments})

@login_required
@user_passes_test(is_super_admin)
def list_departments(request):
    departments = Department.objects.all()
    return render(request, 'list_departments.html', {'departments': departments})

@login_required
@user_passes_test(is_super_admin)
def delete_department(request, pk):
    department = get_object_or_404(Department, pk=pk)
    if request.method == 'POST':
        department.admin.delete()  # Delete the department admin user
        department.delete()
        messages.success(request, 'Department deleted successfully!')
        return redirect('list_departments')
    return render(request, 'delete_department.html', {'department': department})



def teacher_list(request):
    department = Department.objects.get(admin=request.user)
    teachers = Teacher.objects.filter(department=department)
    return render(request, 'teacher_list.html', {'teachers': teachers})

def add_teacher(request):
    try:
        department = Department.objects.get(admin=request.user)
        if request.method == 'POST':
            form = TeacherForm(request.POST, department=department)
            if form.is_valid():
                teacher = form.save(commit=False)
                teacher.department = department
                teacher.save()
                messages.success(request, "Teacher added successfully!")
                return redirect('add_teacher')
            else:
                messages.error(request, "Error adding teacher. Please check the form.")
        else:
            form = TeacherForm(department=department)
        return render(request, 'add_teacher.html', {'form': form, 'teachers': Teacher.objects.filter(department=department)})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to add teachers.')
        return redirect('login')

def delete_teacher(request, teacher_id):
    try:
        department = Department.objects.get(admin=request.user)
        teacher = get_object_or_404(Teacher, id=teacher_id, department=department)
        if request.method == 'POST':
            teacher.delete()
            messages.success(request, "Teacher deleted successfully!")
            return redirect('add_teacher')
        return render(request, 'delete_teacher.html', {'teacher': teacher})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to delete teachers.')
        return redirect('login')

def batch_list(request):
    department = Department.objects.get(admin=request.user)
    batches = Batch.objects.filter(department=department)
    return render(request, 'batch_list.html', {'batches': batches})

def add_batch(request):
    try:
        department = Department.objects.get(admin=request.user)
        if request.method == 'POST':
            form = BatchForm(request.POST, department=department)
            if form.is_valid():
                batch = form.save(commit=False)
                batch.department = department
                batch.save()
                messages.success(request, "Batch added successfully!")
                return redirect('add_batch')
            else:
                messages.error(request, "Error adding batch. Please check the form.")
        else:
            form = BatchForm(department=department)
        return render(request, 'add_batch.html', {'form': form, 'batches': Batch.objects.filter(department=department)})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to add batches.')
        return redirect('login')

def delete_batch(request, batch_id):
    try:
        department = Department.objects.get(admin=request.user)
        batch = get_object_or_404(Batch, id=batch_id, department=department)
        if request.method == 'POST':
            batch.delete()
            messages.success(request, "Batch deleted successfully!")
            return redirect('add_batch')
        return render(request, 'delete_batch.html', {'batch': batch})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to delete batches.')
        return redirect('login')

def division_list(request):
    department = Department.objects.get(admin=request.user)
    divisions = Division.objects.filter(department=department)
    return render(request, 'division_list.html', {'divisions': divisions})

def add_division(request):
    try:
        department = Department.objects.get(admin=request.user)
        if request.method == 'POST':
            form = DivisionForm(request.POST, department=department)
            if form.is_valid():
                division = form.save(commit=False)
                division.department = department
                division.save()
                form.save_m2m()
                messages.success(request, "Division added successfully!")
                return redirect('add_division')
            else:
                messages.error(request, "Error adding division. Please check the form.")
        else:
            form = DivisionForm(department=department)
        return render(request, 'add_division.html', {'form': form, 'divisions': Division.objects.filter(department=department)})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to add divisions.')
        return redirect('login')

def delete_division(request, division_id):
    try:
        department = Department.objects.get(admin=request.user)
        division = get_object_or_404(Division, id=division_id, department=department)
        if request.method == 'POST':
            division.delete()
            messages.success(request, "Division deleted successfully!")
            return redirect('add_division')
        return render(request, 'delete_division.html', {'division': division})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to delete divisions.')
        return redirect('login')



def delete_year(request, year_id):
    try:
        department = Department.objects.get(admin=request.user)
        year = get_object_or_404(Year, id=year_id, department=department)
        if request.method == 'POST':
            year.delete()
            messages.success(request, "Year deleted successfully!")
            return redirect('add_year')
        return render(request, 'delete_year.html', {'year': year})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to delete years.')
        return redirect('login')

def add_year(request):
    try:
        department = Department.objects.get(admin=request.user)
        if request.method == 'POST':
            form = YearForm(request.POST, department=department)
            if form.is_valid():
                year = form.save(commit=False)
                year.department = department
                year.save()
                form.save_m2m()
                messages.success(request, "Year added successfully!")
                return redirect('add_year')
            else:
                messages.error(request, "Error adding year. Please check the form.")
        else:
            form = YearForm(department=department)
        return render(request, 'add_year.html', {'form': form, 'years': Year.objects.filter(department=department)})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to add years.')
        return redirect('login')

def classroom_list(request):
    department = Department.objects.get(admin=request.user)
    classrooms = Classroom.objects.filter(department=department)
    return render(request, 'classroom_list.html', {'classrooms': classrooms})

def add_classroom(request):
    try:
        department = Department.objects.get(admin=request.user)
        if request.method == 'POST':
            form = ClassroomForm(request.POST, department=department)
            if form.is_valid():
                classroom = form.save(commit=False)
                classroom.department = department
                classroom.save()
                messages.success(request, "Classroom added successfully!")
                return redirect('add_classroom')
            else:
                messages.error(request, "Error adding classroom. Please check the form.")
        else:
            form = ClassroomForm(department=department)
        return render(request, 'add_classroom.html', {'form': form, 'classrooms': Classroom.objects.filter(department=department)})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to add classrooms.')
        return redirect('login')

def delete_classroom(request, classroom_id):
    try:
        department = Department.objects.get(admin=request.user)
        classroom = get_object_or_404(Classroom, id=classroom_id, department=department)
        if request.method == 'POST':
            classroom.delete()
            messages.success(request, "Classroom deleted successfully!")
            return redirect('add_classroom')
        return render(request, 'delete_classroom.html', {'classroom': classroom})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to delete classrooms.')
        return redirect('login')


def subject_list(request):
    department = Department.objects.get(admin=request.user)
    subjects = Subject.objects.filter(department=department)
    return render(request, 'subject_list.html', {'subjects': subjects})

def add_subject(request):
    try:
        department = Department.objects.get(admin=request.user)
        if request.method == 'POST':
            form = SubjectForm(request.POST, department=department)
            if form.is_valid():
                subject = form.save(commit=False)
                subject.department = department
                subject.save()
                form.save_m2m()  # Save many-to-many relationships

                # Save batch-teacher mapping in practical_batch_teachers JSONField
                batch_ids = request.POST.getlist('batch_teacher_batch_ids')
                teacher_ids = request.POST.getlist('batch_teacher_teacher_ids')
                mapping = {}
                for batch_id, teacher_id in zip(batch_ids, teacher_ids):
                    if batch_id and teacher_id:
                        mapping[batch_id] = teacher_id
                subject.practical_batch_teachers = mapping
                subject.save()

                messages.success(request, "Subject added successfully!")
                return redirect('add_subject')
            else:
                messages.error(request, "Error adding subject. Please check the form.")
        else:
            form = SubjectForm(department=department)
        # Prepare division->batches and year->divisions mapping
        divisions = Division.objects.filter(department=department)
        years = Year.objects.filter(department=department)
        division_batches = {str(division.id): [batch.id for batch in division.batches.all()] for division in divisions}
        year_divisions = {str(year.id): [division.id for division in year.divisions.all()] for year in years}
        return render(request, 'add_subject.html', {
            'form': form, 
            'subjects': Subject.objects.filter(department=department),
            'divisions': divisions,
            'years': years,
            'division_batches': division_batches,
            'year_divisions': year_divisions,
            'batches': Batch.objects.filter(department=department),
            'teachers': Teacher.objects.filter(department=department),
        })
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to add subjects.')
        return redirect('login')

def view_subject(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    # Prepare batch-wise practical teacher mapping
    batch_teacher_details = []
    if subject.is_practical and subject.practical_batch_teachers:
        for batch_id, teacher_id in subject.practical_batch_teachers.items():
            try:
                batch = Batch.objects.get(id=batch_id)
                teacher = Teacher.objects.get(id=teacher_id)
                batch_teacher_details.append({
                    'batch': batch,
                    'teacher': teacher
                })
            except (Batch.DoesNotExist, Teacher.DoesNotExist):
                continue
    return render(request, 'subject_detail.html', {
        'subject': subject,
        'batch_teacher_details': batch_teacher_details
    })

def edit_subject(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    department = subject.department
    if request.method == 'POST':
        form = SubjectForm(request.POST, instance=subject)
        if form.is_valid():
            subject = form.save()
            # Save batch-teacher mapping in practical_batch_teachers JSONField
            batch_ids = request.POST.getlist('batch_teacher_batch_ids')
            teacher_ids = request.POST.getlist('batch_teacher_teacher_ids')
            mapping = {}
            for batch_id, teacher_id in zip(batch_ids, teacher_ids):
                if batch_id and teacher_id:
                    mapping[batch_id] = teacher_id
            subject.practical_batch_teachers = mapping
            subject.save()
            messages.success(request, "Subject updated successfully!")
            return redirect('add_subject')
        else:
            messages.error(request, "Error updating subject. Please check the form.")
    else:
        form = SubjectForm(instance=subject)
    # Pass division_batches and year_divisions for JS filtering
    divisions = Division.objects.filter(department=department)
    years = Year.objects.filter(department=department)
    division_batches = {str(division.id): [batch.id for batch in division.batches.all()] for division in divisions}
    year_divisions = {str(year.id): [division.id for division in year.divisions.all()] for year in years}
    return render(request, 'edit_subject.html', {
        'form': form, 
        'subject': subject,
        'divisions': divisions,
        'years': years,
        'division_batches': division_batches,
        'year_divisions': year_divisions,
        'batches': Batch.objects.filter(department=department),
        'teachers': Teacher.objects.filter(department=department),
    })

def delete_subject(request, subject_id):
    try:
        department = Department.objects.get(admin=request.user)
        subject = get_object_or_404(Subject, id=subject_id, department=department)
        if request.method == 'POST':
            subject.delete()
            messages.success(request, "Subject deleted successfully!")
            return redirect('add_subject')
        return render(request, 'delete_subject.html', {'subject': subject})
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to delete subjects.')
        return redirect('login')

@login_required
def department_dashboard(request):
    try:
        department = Department.objects.get(admin=request.user)
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to access this page.')
        return redirect('login')
    
    teachers = Teacher.objects.filter(department=department)
    subjects = Subject.objects.filter(department=department)
    classrooms = Classroom.objects.filter(department=department)
    batches = Batch.objects.filter(department=department)
    divisions = Division.objects.filter(department=department)
    years = Year.objects.filter(department=department)
    return render(request, 'admin_dashboard.html', {
        'teachers': teachers,
        'subjects': subjects,
        'classrooms': classrooms,
        'batches': batches,
        'divisions': divisions,
        'years': years,
        
    })
    


@login_required
def view_timetable(request):
    try:
        # Get the logged-in department
        department = Department.objects.get(admin=request.user)
        
        # Get filter parameters
        division_id = request.GET.get('division')
        year_id = request.GET.get('year')
        teacher_id = request.GET.get('teacher')
        classroom_id = request.GET.get('classroom')
        lab_id = request.GET.get('lab')
        batch_id = request.GET.get('batch')
        day = request.GET.get('day')
        
        # Base queryset filtered by department
        timetable_entries = Timetable.objects.filter(subject__department=department)
        
        # Get all data filtered by department
        divisions = Division.objects.filter(department=department)
        years = Year.objects.filter(department=department)
        teachers = Teacher.objects.filter(department=department)
        classrooms = Classroom.objects.filter(department=department, is_lab=False)
        labs = Classroom.objects.filter(department=department, is_lab=True)
        batches = Batch.objects.filter(department=department)
        time_slots = TimeSlot.objects.all().order_by('start_time')
        
        # Apply additional filters
        if division_id:
            timetable_entries = timetable_entries.filter(division_id=division_id)
        
        if year_id:
            timetable_entries = timetable_entries.filter(subject__year_id=year_id)
        
        if teacher_id:
            timetable_entries = timetable_entries.filter(teacher_id=teacher_id)
        
        if classroom_id:
            timetable_entries = timetable_entries.filter(classroom_id=classroom_id)
        
        if lab_id:
            timetable_entries = timetable_entries.filter(lab_id=lab_id)
        
        if batch_id:
            timetable_entries = timetable_entries.filter(batch_id=batch_id)
        
        if day:
            timetable_entries = timetable_entries.filter(day=day)
        
        # Days of the week
        days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
        
        context = {
            'divisions': divisions,
            'years': years,
            'teachers': teachers,
            'classrooms': classrooms,
            'labs': labs,
            'batches': batches,
            'days': days,
            'time_slots': time_slots,
            'timetable_entries': timetable_entries,
            'selected_division': division_id,
            'selected_year': year_id,
            'selected_teacher': teacher_id,
            'selected_classroom': classroom_id,
            'selected_lab': lab_id,
            'selected_batch': batch_id,
            'selected_day': day,
            'department': department
        }
        
        return render(request, 'view_timetable.html', context)
        
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to view timetables.')
        return redirect('login')

# AJAX endpoints for dependent dropdowns
def get_divisions(request):
    department_id = request.GET.get('department')
    divisions = Division.objects.filter(department_id=department_id).values('id', 'name')
    return JsonResponse(list(divisions), safe=False)

def get_years(request):
    department_id = request.GET.get('department')
    years = Year.objects.filter(department_id=department_id).values('id', 'year')
    return JsonResponse(list(years), safe=False)

def get_teachers(request):
    department_id = request.GET.get('department')
    teachers = Teacher.objects.filter(department_id=department_id).values('id', 'name')
    return JsonResponse(list(teachers), safe=False)

def get_classrooms(request):
    department_id = request.GET.get('department')
    classrooms = Classroom.objects.filter(department_id=department_id).values('id', 'name', 'is_lab')
    return JsonResponse(list(classrooms), safe=False)

def get_batches(request):
    department_id = request.GET.get('department')
    batches = Batch.objects.filter(department_id=department_id).values('id', 'name')
    return JsonResponse(list(batches), safe=False)

from collections import defaultdict
from django.shortcuts import render
from .models import Timetable, Classroom, Subject, TimeSlot, Teacher, Division, Batch, Year
import random


def manual_timetable(request):
    try:
        department = Department.objects.get(admin=request.user)
        
        # Get all required data filtered by department
        divisions = Division.objects.filter(department=department)
        years = Year.objects.filter(department=department)
        time_slots = TimeSlot.objects.all().order_by('start_time')
        subjects = Subject.objects.filter(department=department)
        batches = Batch.objects.filter(department=department)
        teachers = Teacher.objects.filter(department=department)
        
        # Get existing timetable entries for this department
        timetable_entries = Timetable.objects.filter(subject__department=department)
        
        # Calculate scheduled hours for each subject and batch
        subject_hours = {}
        batch_teacher_map = {}
        for subject in subjects:
            # Always calculate theory hours if is_theory
            if subject.is_theory:
                scheduled_hours = sum(
                    entry.time_slot.get_duration()
                    for entry in Timetable.objects.filter(subject=subject, batch=None)
                )
                subject_hours.setdefault(subject.id, {})['scheduled'] = scheduled_hours
                subject_hours[subject.id]['total'] = subject.lecture_hours

            # Always calculate practical hours if is_practical
            if subject.is_practical:
                batch_hours = {}
                batch_teacher_map[subject.id] = {}
                for batch in batches:
                    scheduled_hours = sum(
                        entry.time_slot.get_duration()
                        for entry in Timetable.objects.filter(subject=subject, batch=batch)
                    )
                    batch_hours[batch.id] = {
                        'scheduled': scheduled_hours,
                        'total': subject.practical_hours
                    }
                    # Get teacher assignment
                    teacher_id = subject.practical_batch_teachers.get(str(batch.id))
                    if teacher_id:
                        try:
                            teacher = Teacher.objects.get(id=teacher_id)
                            batch_teacher_map[subject.id][str(batch.id)] = {
                                'id': teacher.id,
                                'name': teacher.name
                            }
                        except Teacher.DoesNotExist:
                            batch_teacher_map[subject.id][str(batch.id)] = None
                    else:
                        batch_teacher_map[subject.id][str(batch.id)] = None
                subject_hours.setdefault(subject.id, {})['batch_hours'] = batch_hours
                subject_hours[subject.id]['total'] = subject.practical_hours

        context = {
            'divisions': divisions,
            'years': years,
            'time_slots': time_slots,
            'subjects': subjects,
            'batches': batches,
            'DAYS': ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday'],
            'subject_hours': subject_hours,
            'timetable_entries': timetable_entries,
            'department': department,
            'teachers': teachers,
            'batch_teacher_map': batch_teacher_map,  # Add the batch-teacher mapping to context
        }
        
        return render(request, 'manual_timetable.html', context)
    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to access this page.')
        return redirect('login')

def delete_all_timetable_entries(request):
    Timetable.objects.all().delete()
    messages.success(request, "All timetable entries have been deleted successfully!")
    return redirect('manual_timetable')



@login_required
def save_timetable_entry(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        day = data['day']
        time_slot = TimeSlot.objects.get(id=data['time_slot_id'])
        division = Division.objects.get(id=data['division_id'])
        year = Year.objects.get(id=data['year_id'])
        subject = Subject.objects.get(id=data['subject_id'])
        is_practical = data.get('is_practical', False)

        new_start = time_slot.start_time
        new_end = time_slot.end_time

        year_division_ids = year.divisions.values_list('id', flat=True)

        # ✅ Filter existing timetable entries for the SAME division and SAME year only
        overlapping_entries = Timetable.objects.filter(
                day=data['day'],
                division=division,
                subject__year=year
            )
        # ✅ Check for overlapping time slot
        for entry in overlapping_entries:
            existing_start = entry.time_slot.start_time
            existing_end = entry.time_slot.end_time

            if new_start < existing_end and new_end > existing_start:
                from_time = existing_start.strftime("%I:%M %p") if existing_start else "N/A"
                to_time = existing_end.strftime("%I:%M %p") if existing_end else "N/A"

                if not is_practical:
                    return JsonResponse({
                        'success': False,
                        'message': f'This time slot overlaps with an existing {"practical" if entry.lab else "lecture"} ({entry.subject.name}) from {from_time} to {to_time}'
                    })
                else:
                    if not entry.lab:  # If existing is a lecture, block practical
                        return JsonResponse({
                            'success': False,
                            'message': f'This time slot overlaps with an existing lecture ({entry.subject.name}) from {from_time} to {to_time}'
                        })

        # ✅ No conflicts found, proceed with saving
        if is_practical:
            batch = Batch.objects.get(id=data['batch_id'])
            lab = subject.lab

            teacher_id = subject.practical_batch_teachers.get(str(batch.id))
            if not teacher_id:
                return JsonResponse({
                    'success': False,
                    'message': f'No teacher assigned for batch {batch.name}'
                })

            teacher = Teacher.objects.get(id=teacher_id)

            entry = Timetable.objects.create(
                subject=subject,
                time_slot=time_slot,
                division=division,
                day=day,
                teacher=teacher,
                batch=batch,
                lab=lab,
                classroom=None
            )

            scheduled_hours = sum(
                e.time_slot.get_duration()
                for e in Timetable.objects.filter(subject=subject, batch=batch)
            )

            return JsonResponse({
                'success': True,
                'message': 'Entry added successfully',
                'entry_id': entry.id,
                'teacher_name': teacher.name,
                'lab_name': lab.name if lab else None,
                'batch_name': batch.name,
                'hours': {
                    'scheduled': scheduled_hours,
                    'total': subject.practical_hours
                }
            })

        else:
            if not subject.theory_teacher:
                return JsonResponse({
                    'success': False,
                    'message': f'No theory teacher assigned for subject {subject.name}'
                })

            teacher = subject.theory_teacher

            entry = Timetable.objects.create(
                subject=subject,
                time_slot=time_slot,
                division=division,
                day=day,
                teacher=teacher,
                batch=None,
                lab=None,
                classroom=year.classroom
            )

            scheduled_hours = sum(
                e.time_slot.get_duration()
                for e in Timetable.objects.filter(subject=subject)
            )

            return JsonResponse({
                'success': True,
                'message': 'Entry added successfully',
                'entry_id': entry.id,
                'teacher_name': teacher.name,
                'classroom_name': year.classroom.name if year.classroom else None,
                'hours': {
                    'scheduled': scheduled_hours,
                    'total': subject.lecture_hours
                }
            })

    return JsonResponse({
        'success': False,
        'message': 'Invalid request method'
    })


@login_required
def view_department_timetable(request):
    if not request.user.is_superuser:
        messages.error(request, "You don't have permission to access this page.")
        return redirect('home')
    
    department_id = request.GET.get('department')
    if not department_id:
        messages.error(request, "Please select a department.")
        return redirect('main_admin_dashboard')
    
    try:
        department = Department.objects.get(id=department_id)
    except Department.DoesNotExist:
        messages.error(request, "Department not found.")
        return redirect('main_admin_dashboard')
    
    # Get all divisions for this department
    divisions = Division.objects.filter(department=department)
    
    # Get all time slots
    time_slots = TimeSlot.objects.all().order_by('start_time')
    
    # Get all timetable entries for this department
    timetable_entries = Timetable.objects.filter(
        division__department=department
    ).select_related(
        'subject',
        'teacher',
        'classroom',
        'lab',
        'batch',
        'division',
        'subject__year'
    )
    
    context = {
        'department': department,
        'divisions': divisions,
        'time_slots': time_slots,
        'timetable_entries': timetable_entries,
        'DAYS': ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday'],
    }
    
    return render(request, 'view_department_timetable.html', context)

@login_required
def view_meta_timetable(request):
    if not request.user.is_superuser:
        messages.error(request, "You don't have permission to access this page.")
        return redirect('home')
    
    # Get all departments
    departments = Department.objects.all()
    
    # Get all time slots
    time_slots = TimeSlot.objects.all().order_by('start_time')
    
    # Get all timetable entries
    timetable_entries = Timetable.objects.all().select_related(
        'subject',
        'teacher',
        'classroom',
        'lab',
        'batch',
        'division',
        'subject__year',
        'division__department'
    )
    
    # Get filter parameters from request
    department_id = request.GET.get('department')
    division_id = request.GET.get('division')
    year_id = request.GET.get('year')
    teacher_id = request.GET.get('teacher')
    classroom_id = request.GET.get('classroom')
    lab_id = request.GET.get('lab')
    batch_id = request.GET.get('batch')
    day = request.GET.get('day')
    
    # Apply filters
    if department_id:
        timetable_entries = timetable_entries.filter(division__department_id=department_id)
    if division_id:
        timetable_entries = timetable_entries.filter(division_id=division_id)
    if year_id:
        timetable_entries = timetable_entries.filter(subject__year_id=year_id)
    if teacher_id:
        timetable_entries = timetable_entries.filter(teacher_id=teacher_id)
    if classroom_id:
        timetable_entries = timetable_entries.filter(classroom_id=classroom_id)
    if lab_id:
        timetable_entries = timetable_entries.filter(lab_id=lab_id)
    if batch_id:
        timetable_entries = timetable_entries.filter(batch_id=batch_id)
    if day:
        timetable_entries = timetable_entries.filter(day=day)
    
    # Get all divisions, years, teachers, classrooms, labs, and batches for the selected department
    divisions = Division.objects.all()
    years = Year.objects.all()
    teachers = Teacher.objects.all()
    classrooms = Classroom.objects.filter(is_lab=False)
    labs = Classroom.objects.filter(is_lab=True)
    batches = Batch.objects.all()
    
    # Days of the week
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    
    context = {
        'departments': departments,
        'divisions': divisions,
        'years': years,
        'teachers': teachers,
        'classrooms': classrooms,
        'labs': labs,
        'batches': batches,
        'days': days,
        'time_slots': time_slots,
        'timetable_entries': timetable_entries,
        'selected_department': department_id,
        'selected_division': division_id,
        'selected_year': year_id,
        'selected_teacher': teacher_id,
        'selected_classroom': classroom_id,
        'selected_lab': lab_id,
        'selected_batch': batch_id,
        'selected_day': day,
    }
    
    return render(request, 'view_meta_timetable.html', context)

@login_required
def select_department_timetable(request):
    if not request.user.is_superuser:
        messages.error(request, "You don't have permission to access this page.")
        return redirect('home')
    
    departments = Department.objects.all()
    return render(request, 'select_department_timetable.html', {'departments': departments})

@login_required
def delete_timetable_entry(request, entry_id):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Invalid request method'})
    
    try:
        entry = Timetable.objects.get(id=entry_id)
        # Check if the user has permission to delete this entry
        if not request.user.is_superuser and entry.division.department.admin != request.user:
            return JsonResponse({'success': False, 'error': 'Permission denied'})
        
        entry.delete()
        return JsonResponse({'success': True})
    except Timetable.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Entry not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
def download_timetable(request, format):
    try:
        # Get filter parameters from request
        department_id = request.GET.get('department')
        division_id = request.GET.get('division')
        year_id = request.GET.get('year')
        teacher_id = request.GET.get('teacher')
        classroom_id = request.GET.get('classroom')
        lab_id = request.GET.get('lab')
        batch_id = request.GET.get('batch')
        day = request.GET.get('day')

        # Determine department context
        if request.user.is_superuser:
            # For superusers viewing meta timetable
            if department_id:
                department = Department.objects.get(id=department_id)
                timetable_entries = Timetable.objects.filter(division__department=department)
            else:
                department = None
                timetable_entries = Timetable.objects.all()
        else:
            # For department admins
            department = Department.objects.get(admin=request.user)
            timetable_entries = Timetable.objects.filter(subject__department=department)

        # Apply filters
        if division_id:
            timetable_entries = timetable_entries.filter(division_id=division_id)
        if year_id:
            timetable_entries = timetable_entries.filter(subject__year_id=year_id)
        if teacher_id:
            timetable_entries = timetable_entries.filter(teacher_id=teacher_id)
        if classroom_id:
            timetable_entries = timetable_entries.filter(classroom_id=classroom_id)
        if lab_id:
            timetable_entries = timetable_entries.filter(lab_id=lab_id)
        if batch_id:
            timetable_entries = timetable_entries.filter(batch_id=batch_id)
        if day:
            timetable_entries = timetable_entries.filter(day=day)
        
        time_slots = TimeSlot.objects.all().order_by('start_time')
        days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']

        # Create filter info text
        filter_info = []
        if division_id:
            division = Division.objects.get(id=division_id)
            filter_info.append(f"Division: {division.name}")
        if year_id:
            year = Year.objects.get(id=year_id)
            filter_info.append(f"Year: {year.year}")
        if teacher_id:
            teacher = Teacher.objects.get(id=teacher_id)
            filter_info.append(f"Teacher: {teacher.name}")
        if classroom_id:
            classroom = Classroom.objects.get(id=classroom_id)
            filter_info.append(f"Classroom: {classroom.name}")
        if lab_id:
            lab = Classroom.objects.get(id=lab_id)
            filter_info.append(f"Lab: {lab.name}")
        if batch_id:
            batch = Batch.objects.get(id=batch_id)
            filter_info.append(f"Batch: {batch.name}")
        if day:
            filter_info.append(f"Day: {day}")

        generated_on = datetime.now().strftime('%d-%m-%Y %I:%M %p')

        department_label = department.name if department else "All Departments"

        # For PDF/Word/Excel, add department info if department is None (meta timetable)
        show_dept = department is None

        if format == 'pdf':
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
            elements = []

            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                spaceAfter=10,
                alignment=1,
                textColor=colors.HexColor('#2c3e50')
            )
            college_style = ParagraphStyle(
                'CollegeName',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=5,
                alignment=1,
                textColor=colors.HexColor('#2c3e50')
            )
            info_style = ParagraphStyle(
                'Info',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=5,
                alignment=1,
                textColor=colors.HexColor('#666666')
            )
            cell_style = ParagraphStyle(
                'Cell',
                parent=styles['Normal'],
                fontSize=9,
                alignment=TA_CENTER,
                leading=12,
                spaceAfter=2,
            )
            elements.append(Paragraph("GOVERNMENT POLYTECHNIC, NASHIK", college_style))
            elements.append(Paragraph(f"Timetable - {department_label}", title_style))
            if filter_info:
                elements.append(Paragraph("Filters: " + ", ".join(filter_info), info_style))
            elements.append(Paragraph(f"Generated on: {generated_on}", info_style))
            elements.append(Spacer(1, 0.15*inch))

            # Table header
            table_data = [
                [Paragraph('<b>Day</b>', cell_style)] + [Paragraph(f"<b>{time_filter(slot.start_time, 'g:i A')} - {time_filter(slot.end_time, 'g:i A')}</b>", cell_style) for slot in time_slots]
            ]

            for day in days:
                row = [Paragraph(f'<b>{day}</b>', cell_style)]
                for slot in time_slots:
                    entries = timetable_entries.filter(day=day, time_slot=slot)
                    cell_content = []
                    for entry in entries:
                        lines = [
                            f'{entry.subject.name} - {entry.teacher.name}',
                            f'{entry.lab.name if entry.lab else entry.classroom.name}',
                            f'{"Batch" if entry.batch else "Division"}: {entry.batch.name if entry.batch else entry.division.name}'
                        ]
                        if show_dept:
                            lines.append(f'Department: {entry.division.department.name}')
                        cell_content.append('<br/>'.join(lines))
                    if cell_content:
                        para = Paragraph('<br/><br/>'.join(cell_content), cell_style)
                        row.append(para)
                    else:
                        row.append(Paragraph('', cell_style))
                table_data.append(row)

            # Calculate dynamic column widths for PDF
            total_width = 11.0 * inch  # Letter landscape width
            min_col_width = 1.2 * inch
            n_slots = len(time_slots)
            # Calculate width for each slot column so all fit, but not less than min_col_width
            slot_col_width = max((total_width - min_col_width) / n_slots, min_col_width)
            if slot_col_width * n_slots + min_col_width > total_width:
                slot_col_width = (total_width - min_col_width) / n_slots
            col_widths = [min_col_width] + [slot_col_width] * n_slots
            table = Table(table_data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a90e2')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(table)
            doc.build(elements)
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="timetable_{department_label}.pdf"'
            return response

        elif format == 'docx':
            doc = Document()
            # Set landscape orientation
            section = doc.sections[-1]
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            # Title: College name
            college_para = doc.add_paragraph()
            college_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            college_run = college_para.add_run("GOVERNMENT POLYTECHNIC, NASHIK")
            college_run.bold = True
            college_run.font.size = Pt(16)
            # Title: Department name
            main_para = doc.add_paragraph()
            main_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            main_run = main_para.add_run(f"Timetable - {department_label}")
            main_run.bold = True
            main_run.font.size = Pt(14)
            # Filters
            if filter_info:
                filter_para = doc.add_paragraph()
                filter_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                filter_run = filter_para.add_run("Filters: " + ", ".join(filter_info))
                filter_run.font.size = Pt(10)
                filter_run.font.color.rgb = RGBColor(102, 102, 102)
            # Generated on
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info_para.add_run(f"Generated on: {generated_on}")
            info_run.font.size = Pt(10)
            info_run.font.color.rgb = RGBColor(102, 102, 102)
            # Table
            table = doc.add_table(rows=1, cols=len(time_slots) + 1)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Day'
            for i, slot in enumerate(time_slots, 1):
                header_cells[i].text = f"{time_filter(slot.start_time, 'g:i A')} - {time_filter(slot.end_time, 'g:i A')}"
            # Style header row
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '4a90e2')
                tcPr.append(shd)
                cell.vertical_alignment = 1
            # Data rows
            for day in days:
                row_cells = table.add_row().cells
                row_cells[0].text = day
                for i, slot in enumerate(time_slots, 1):
                    entries = timetable_entries.filter(day=day, time_slot=slot)
                    cell_content = []
                    for entry in entries:
                        lines = [
                            f"{entry.subject.name} - {entry.teacher.name}",
                            f"{entry.lab.name if entry.lab else entry.classroom.name}",
                            f"{'Batch' if entry.batch else 'Division'}: {entry.batch.name if entry.batch else entry.division.name}"
                        ]
                        if show_dept:
                            lines.append(f"Department: {entry.division.department.name}")
                        cell_content.append('\n'.join(lines))
                    row_cells[i].text = '\n\n'.join(cell_content) if cell_content else ''
                    # Center align all cell text
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                # Center align day cell
                for paragraph in row_cells[0].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            # Set column widths for landscape
            word_total_width = 12.0  # inches, for landscape
            min_col_width = 1.0
            n_slots = len(time_slots)
            slot_col_width = max((word_total_width - min_col_width) / n_slots, min_col_width)
            if slot_col_width * n_slots + min_col_width > word_total_width:
                slot_col_width = (word_total_width - min_col_width) / n_slots
            col_widths = [min_col_width] + [slot_col_width] * n_slots
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    cell.width = Inches(col_widths[idx])
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn2('w:w'), str(int(col_widths[idx]*1440)))
                    tcW.set(qn2('w:type'), 'dxa')
                    tcPr.append(tcW)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="timetable_{department_label}.docx"'
            return response

        elif format == 'xlsx':
            excel_data = []
            for day in days:
                for slot in time_slots:
                    entries = timetable_entries.filter(day=day, time_slot=slot)
                    for entry in entries:
                        row = {
                            'Day': day,
                            'Time': f"{time_filter(slot.start_time, 'g:i A')} - {time_filter(slot.end_time, 'g:i A')}",
                            'Subject - Teacher': f"{entry.subject.name} - {entry.teacher.name}",
                            'Room': entry.lab.name if entry.lab else entry.classroom.name,
                            'Batch/Division': entry.batch.name if entry.batch else entry.division.name
                        }
                        if show_dept:
                            row['Department'] = entry.division.department.name
                        excel_data.append(row)
            excel_df = pd.DataFrame(excel_data)
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                # Write timetable data starting from row 5 (leave space for title, filters, date)
                excel_df.to_excel(writer, index=False, sheet_name='Timetable', startrow=4)
                worksheet = writer.sheets['Timetable']
                # Write title, filters, and generated date in the first rows
                worksheet.cell(row=1, column=1, value="GOVERNMENT POLYTECHNIC, NASHIK")
                worksheet.cell(row=2, column=1, value=f"Timetable - {department_label}")
                if filter_info:
                    worksheet.cell(row=3, column=1, value="Filters: " + ', '.join(filter_info))
                worksheet.cell(row=4, column=1, value=f"Generated on: {generated_on}")
                # Style header row (row 5)
                header_fill = openpyxl.styles.PatternFill(start_color='4a90e2', end_color='4a90e2', fill_type='solid')
                header_font = openpyxl.styles.Font(color='FFFFFF', bold=True)
                for cell in worksheet[5]:
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center')
                # Center-align all data cells
                for row in worksheet.iter_rows(min_row=6, max_row=worksheet.max_row, min_col=1, max_col=worksheet.max_column):
                    for cell in row:
                        cell.alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center', wrap_text=True)
                # Auto-fit column widths
                for column in worksheet.columns:
                    max_length = 0
                    column = [cell for cell in column]
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = (max_length + 2)
                    worksheet.column_dimensions[openpyxl.utils.get_column_letter(column[0].column)].width = adjusted_width
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Disposition'] = f'attachment; filename="timetable_{department_label}.xlsx"'
            return response
        else:
            messages.error(request, 'Invalid format specified')
            return redirect('view_timetable')

    except Department.DoesNotExist:
        messages.error(request, 'You are not authorized to download the timetable.')
        return redirect('login')



